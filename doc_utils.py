from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Cm, Pt, RGBColor
import datetime


def create_doc(landscape=False, my=None, mx=None):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    section = doc.sections[-1]
    if landscape:
        old_width, old_height = section.page_width, section.page_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = old_width
        section.page_width = old_height

    if my:
        section.top_margin = Cm(my)
        section.bottom_margin = Cm(my)

    if mx:
        section.left_margin = Cm(mx)
        section.right_margin = Cm(mx)
    
    doc.core_properties.author = 'TBTA'
    doc.core_properties.created = datetime.datetime.today()
    doc.core_properties.modified = datetime.datetime.today()
    return doc


def format_run(paragraph, run_data):
    if isinstance(run_data, dict):
        # run_data is expected in the format { text, bold?, red?, highlight?, size? }
        run = paragraph.add_run(text=run_data['text'])
        if 'bold' in run_data and run_data['bold']:
            run.font.bold = True
        if 'red' in run_data and run_data['red']:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        if 'highlight' in run_data and run_data['highlight']:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if 'size' in run_data:
            run.font.size = Pt(run_data['size'])
    else:
        paragraph.add_run(text=str(run_data))


def format_paragraph(paragraph, text_data):
    if isinstance(text_data, list):
        for run_data in text_data:
            format_run(paragraph, run_data)
    else:
        format_run(paragraph, text_data)


def add_paragraph(doc, text='', formatting={}):
    paragraph = doc.add_paragraph()
    if text:
        format_paragraph(paragraph, text)
    if 'space_after' in formatting:
        paragraph.paragraph_format.space_after = Pt(formatting['space_after'])
    if 'center' in formatting and formatting['center']:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_table(doc, rows, col_widths, caption=None):
    if not len(rows):
        return None
    
    if caption:
        add_paragraph(doc, caption, formatting={ 'center': True, 'space_after': 0 })

    table = doc.add_table(rows=0, cols=len(col_widths), style='Table Grid')

    for row_data in rows:
        # Note: this is MUCH faster than accessing the row cells each time within the cell loop
        row_cells = table.add_row().cells

        for col_num, col_width in enumerate(col_widths):
            cell = row_cells[col_num]
            if col_num < len(row_data):
                format_paragraph(cell.paragraphs[0], row_data[col_num])

            # Each cell needs its width set individually for some reason
            cell.width = Cm(col_width)

    return table