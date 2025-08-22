import sys
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Cm


# Parameter Name Constants
PARAM_INPUT_PATH = 'input_path'
PARAM_OUTPUT_PATH = 'output_path'
PARAM_NOTES_COLUMN = 'add_notes_column'
PARAM_PASSAGE = 'passage'

# Concept Info Fields
CONCEPT_WORD = 'word'
CONCEPT_GLOSS = 'gloss'
CONCEPT_VERSE_REF = 'verse_ref'
CONCEPT_VERSE_TEXT = 'verse_text'
CONCEPT_OCCURRENCES = 'verse_occurrences'
CONCEPT_SAMPLE = 'sample'

# Semantic Categories
CATEGORY_PROPER = 'Proper Name'
CATEGORY_NOUN = 'Noun'
CATEGORY_VERB = 'Verb'
CATEGORY_ADJECTIVE = 'Adjective'
CATEGORY_ADVERB = 'Adverb'
CATEGORY_ADPOSITION = 'Adposition'
CATEGORY_CONJUNCTION = 'Conjunction'
CATEGORY_PARTICLE = 'Particle'
CATEGORY_PHRASAL = 'Phrasal'

# Table Header Names
HEADER_GLOSS = 'Glosses'
HEADER_OCCURRENCE = 'Verses'
HEADER_SAMPLE = 'Target Sentences'
HEADER_TARGET_WORD = 'Target Words'
HEADER_TARGET_GLOSS = 'Target Glosses'
HEADER_NOTES = 'Notes'


def get_params():
    if len(sys.argv) < 2:
        print('Please specify a .txt file to import')
        return None

    file_name = sys.argv[-1]
    if file_name.startswith('-'):
        print('File name must be the last argument')
        return None

    file_path = Path(file_name).with_suffix('.txt')
    if not file_path.exists():
        print('Specified File does not exist...')
        return None

    return {
        PARAM_INPUT_PATH: file_path,
        PARAM_OUTPUT_PATH: file_path.with_name(f'Lexicon - {file_path.stem}.docx'),
        PARAM_NOTES_COLUMN: '-N' in sys.argv or '-n' in sys.argv,
    }


def import_concepts(params):
    CONCEPT_REGEX = re.compile(r'^Concept \((?P<category>[a-zA-Z]+)\): (?P<word>[.a-zA-Z0-9- ]+?-[A-Z])(?:  \'(?P<gloss>.+?)\')?$')
    VERSE_REGEX = re.compile(r'^Verse: (?P<ref>[.a-zA-Z0-9- ]+:\d+) ?(?P<text>.*)$')
    GLOSS_REPLACE_REGEX = re.compile(r'\((LDV|simple|inexplicable|proper name|universal primitive)\) ')
    categories = {}

    def add_concept_to_category(concept, category):
        if category == CATEGORY_NOUN and concept[CONCEPT_WORD][0].isupper():
            category = CATEGORY_PROPER
        if category not in categories:
            categories[category] = []
        categories[category].append(concept)
        return category

    path = params[PARAM_INPUT_PATH]
    with path.open() as f:
        concept = None
        category = ''
        for line_num, line in enumerate(f):
            if line.startswith('Concept'):
                concept_match = CONCEPT_REGEX.match(line)
                if not concept_match:
                    print('Unexpected format for Concept on line ' + str(line_num))
                    continue

                concept = {
                    CONCEPT_WORD: concept_match['word'],
                    CONCEPT_GLOSS: GLOSS_REPLACE_REGEX.sub('', concept_match['gloss']) if concept_match['gloss'] else ''
                }
                category = add_concept_to_category(concept, concept_match['category'])

            elif line.startswith('Sample Sentence'):
                concept[CONCEPT_SAMPLE] = line[len('Sample Sentence: '):].strip()

            elif line.startswith('Verse'):
                if category == CATEGORY_PROPER:
                    continue
                verse_match = VERSE_REGEX.match(line)
                concept[CONCEPT_VERSE_REF] = verse_match['ref']
                concept[CONCEPT_VERSE_TEXT] = verse_match['text']
                concept[CONCEPT_OCCURRENCES] = extract_verse_occurrences(concept[CONCEPT_WORD], verse_match['text'])

            elif line.startswith('Current Passage'):
                # This only appears once at the top of the text file
                params[PARAM_PASSAGE] = line[len('Current Passage: '):].strip()

    # Only for debug purposes
    # for k,v in categories.items():
    #     print(f'Retrieved {len(v)} {k} unlinked concepts from "{path}"')
    return categories


SENTENCE_REGEX = re.compile(r'([^.?!]+[.?!]\S*) ?')
def extract_verse_occurrences(word, text):
    # Remove the sense -X from the word
    dash_idx = word.rindex('-')
    word = word[:dash_idx]
    
    occurrences = []

    # The verse might have an unrecognizable form of the word, the word
    # might not be present at all due to restructuring. 
    word_regex = r'\b(' + re.escape(word) + r'(?:s|es|ed|d)?)\b'
    for sentence in SENTENCE_REGEX.findall(text):
        word_match = re.search(word_regex, sentence, re.IGNORECASE)
        if word_match:
            occurrences.append({
                'sentence': sentence,
                'location': word_match.span(1),
            })

    return occurrences


def export_document(categories, params):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    # Set orientation to landscape
    section = doc.sections[-1]
    old_width, old_height = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = old_width
    section.page_width = old_height

    # Set the margins to Normal (2.54cm for each)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Add the passage as a heading
    passage_paragraph = doc.add_paragraph(params[PARAM_PASSAGE])
    passage_run = passage_paragraph.runs[0]
    passage_run.bold = True
    passage_run.font.size = Pt(14)

    # Create the tables
    table_order = [
        CATEGORY_PROPER,
        CATEGORY_NOUN,
        CATEGORY_VERB,
        CATEGORY_ADJECTIVE,
        CATEGORY_ADVERB,
        CATEGORY_ADPOSITION,
        CATEGORY_CONJUNCTION,
        CATEGORY_PARTICLE,
        CATEGORY_PHRASAL,
    ]
    ordered_categories = sorted(categories.items(), key=lambda kv: table_order.index(kv[0]))
    add_notes_column = params[PARAM_NOTES_COLUMN]
    for idx, (category, concepts) in enumerate(ordered_categories):
        create_table(category, concepts, idx+1, doc, add_notes_column)

    try:
        doc.save(str(params[PARAM_OUTPUT_PATH]))
        return True
    except PermissionError:
        err_text = f'"{params[PARAM_OUTPUT_PATH].name}" is currently open. Please close and try again.'
        print("Error: " + err_text)
        import ctypes  
        ctypes.windll.user32.MessageBoxW(0, err_text, "Error Creating Word Document", 0 + 16)
        return False


def create_table(category, concepts, table_num, doc, add_notes_column):
    # Figure out the column names and widths
    if category == CATEGORY_PROPER:
        col_names = [f'Nouns: {category}s', HEADER_GLOSS, HEADER_TARGET_WORD]
        col_widths = [Cm(5)] * 3
    elif add_notes_column:
        col_names = [category + 's', HEADER_GLOSS, HEADER_OCCURRENCE, HEADER_SAMPLE, HEADER_TARGET_WORD, HEADER_TARGET_GLOSS, HEADER_NOTES]
        col_widths = [Cm(2.5), Cm(3), Cm(6), Cm(3), Cm(3), Cm(3), Cm(3)]
    else:
        col_names = [category + 's', HEADER_GLOSS, HEADER_OCCURRENCE, HEADER_SAMPLE, HEADER_TARGET_WORD, HEADER_TARGET_GLOSS]
        col_widths = [Cm(2.7), Cm(3.2), Cm(6.3), Cm(4), Cm(3.4), Cm(3.4)]

    # Add the caption
    if table_num > 1:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    caption = doc.add_paragraph(f'Table {table_num}. {category}s')
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)

    # Create the table
    table = doc.add_table(rows=1, cols=len(col_names), style='Table Grid')

    # Add the header
    header_cols = table.row_cells(0)
    for idx, text in enumerate(col_names):
        header_cols[idx].text = text
        if text == HEADER_TARGET_WORD:
            header_cols[idx].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    # TODO do we need to alphabetize the concepts?
    for concept in concepts:
        row_cells = table.add_row().cells
        row_cells[0].text = concept[CONCEPT_WORD]
        row_cells[1].text = concept[CONCEPT_GLOSS]

        if CONCEPT_OCCURRENCES in concept:
            add_verse_sentences(concept, row_cells[2])
        if CONCEPT_SAMPLE in concept:
            add_sample_sentences(concept, row_cells[3])

    # Each cell width needs to be set individually
    for idx, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[idx]


def add_verse_sentences(concept, table_cell):
    paragraph = table_cell.paragraphs[0]

    if concept[CONCEPT_OCCURRENCES]:
        paragraph.add_run(text=concept[CONCEPT_VERSE_REF])

        # Show each occurrence of the word in bold
        for occurrence in concept[CONCEPT_OCCURRENCES]:
            sentence = occurrence['sentence']
            start, end = occurrence['location']
            paragraph.add_run(text=' ' + sentence[:start])
            paragraph.add_run(text=sentence[start:end]).bold = True
            paragraph.add_run(text=sentence[end:])
    else:
        # Show the whole verse and highlight the text so the user knows to attend to it
        run_text = concept[CONCEPT_VERSE_REF] + ' ' + concept[CONCEPT_VERSE_TEXT]
        run_font = paragraph.add_run(text=run_text).font
        run_font.highlight_color = WD_COLOR_INDEX.YELLOW

    for run in paragraph.runs:
        run.font.size = Pt(10)


def add_sample_sentences(concept, table_cell):
    # Start with the sample sentence itself with the separating bar
    table_cell.text = concept[CONCEPT_SAMPLE] + ' | '
    paragraph = table_cell.paragraphs[0]

    # Add the place where the translator needs to translate
    run_font = paragraph.add_run(text='Translation here.').font
    run_font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Set the font size
    for run in paragraph.runs:
        run.font.size = Pt(10)


if __name__ == "__main__":
    params = get_params()
    if params:
        concepts = import_concepts(params)
        if export_document(concepts, params):
            params[PARAM_INPUT_PATH].unlink()   # delete the original text file
            print(f'Successfully exported "{params[PARAM_OUTPUT_PATH]}"')
