# This script takes the exported text file from TBTA and arranges the verses
# into a Word document table that can then be sent to the MTT. It separates the
# English and the target language, making sure the corresponding verses and
# sentences are aligned properly.
#
# The txt file to import the verses from must be specified as the last argument.
# Including the argument -s will separate each sentence into its own line.
#   By default each whole verse gets its own row
# Including the argument -n will add a blank 'Notes' column on the right.
#   By default it is excluded.
#
# When splitting sentences, there may be some misalignment between the two
# langauges due to some sentences being combined in one language but not the other,
# or due to the presence of an implicit sentence. A blank line will alert the
# user that there is misalignment which will have to be dealt with manually.
# However, this misalignment will not extend beyond the verse in question.

import sys
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm


class VerseInfo:
    def __init__(self):
        self.ref = ''
        self.languages = []
        self.previous_was_title = False

    def append_title(self, text):
        # A title is on the line before the main verse, even before
        # the verse reference.
        self.previous_was_title = True
        self.languages.append(text.strip())
    
    def append_verse(self, text):
        if self.previous_was_title:
            self.languages[-1] += ' ' + text.strip()
        else:
            self.languages.append(text.strip())
        self.previous_was_title = False

    def append_footnote(self, text):
        self.languages[-1] += ' ' + text.strip()
        self.previous_was_title = False

    def __repr__(self) -> str:
        return f'ref: {self.ref}; prev_was_title: {self.previous_was_title}; text: {self.languages}'

    
# Parameter Name constants
PARAM_INPUT_PATH = 'input_path'
PARAM_OUTPUT_PATH = 'output_path'
PARAM_FOOTNOTE = 'footnote'
PARAM_SPLIT_SENTENCES = 'split_sentences'
PARAM_NOTES_COLUMN = 'add_notes_column'
PARAM_PUB_REFS = 'publishable_refs'

def get_params():
    # usage is: tbta_export_to_word.exe -s -n -p "text_file.txt" "footnote:"
    # The text file path and footnote word are required
    if len([a for a in sys.argv if not a.startswith('-')]) < 3:
        show_error('Please specify a .txt file to import and the Target word for Footnote')
        return None

    file_name = sys.argv[-2]
    file_path = Path(file_name).with_suffix('.txt')
    if not file_path.exists():
        show_error(f'Specified File "{file_name}" does not exist...')
        return None

    return {
        PARAM_INPUT_PATH: file_path,
        PARAM_OUTPUT_PATH: file_path.with_name(f'{file_path.stem}.docx'),
        PARAM_FOOTNOTE: sys.argv[-1],
        PARAM_SPLIT_SENTENCES: '-S' in sys.argv or '-s' in sys.argv,
        PARAM_NOTES_COLUMN: '-N' in sys.argv or '-n' in sys.argv,
        PARAM_PUB_REFS: '-P' in sys.argv or '-p' in sys.argv,
    }


def get_language_info(file_iter):
    # The first line has the language info.
    # This handles the following cases:
    #   Target-Churched Adults and English-Churched Adults and Third-Churched Adults
    #   Target and English-Churched Adults
    #   Target-Churched Adults
    #   Target
    line = next(file_iter)
    languages = [s.split('-')[0] for s in line.strip().split(' and ')]

    # The third line has the book name for each language
    next(file_iter)
    line = next(file_iter)
    book_names = re.findall(r'\s*(.+?)\s*(?:$|-)', line)

    def create_info(i, is_target=False):
        if is_target:
            return { 'name': languages.pop(i), 'book_name': book_names[0] }
        else:
            return { 'name': languages.pop(i) }

    # If there is only one language, it is always the target language (which might even be English)
    if len(languages) == 1:
        return [create_info(0, True)]
    
    # In the file the order is Target (English) (Other)
    # Reorder so it's (English) Target (Other) which is the order the verses appear in
    language_info = []
    if 'English' in languages:
        language_info.append(create_info(1, False))
    language_info.append(create_info(0, True))
    if len(languages) > 0:
        language_info.append(create_info(0))

    return language_info


def import_file(encoding, params):
    with params[PARAM_INPUT_PATH].open(encoding=encoding, newline='\r\n') as file:
        file_iter = iter(file)
        language_info = get_language_info(file_iter)
        print(language_info)

        if language_info is None:
            return (None, None)
        if (params[PARAM_PUB_REFS] or 
                (len(language_info) == 1 and not params[PARAM_SPLIT_SENTENCES] and not params[PARAM_NOTES_COLUMN])):
            target_info = [lang for lang in language_info if 'book_name' in lang][0]
            title_line = target_info['name'] + ' - ' + target_info['book_name']
            lines = [title_line] + [line.strip() for line in file_iter]
            return (lines, None)
        else:
            table_verses = import_verses_for_table(file_iter, language_info, params)
            params['languages'] = [x['name'] for x in language_info]
            return (None, table_verses)


def import_verses_for_table(file_iter, language_info, params):
    footnote_word = params[PARAM_FOOTNOTE]

    verses = []
    current_verse = VerseInfo()

    has_english = any([x['name'] == 'English' for x in language_info])
    def target_index():
        return 1 if has_english and len(language_info) > 1 else 0

    def is_target():
        return len(current_verse.languages) == (target_index() + 1)

    def check_verse_boundary():
        nonlocal current_verse
        if len(current_verse.languages) == len(language_info) and not current_verse.previous_was_title:
            verses.append(current_verse)
            current_verse = VerseInfo()

    # CHAPTER_REGEX = re.compile(r'([^\d]+ \d+)\n')
    VERSE_LINE_REGEX = re.compile(r'^(.+? \d+:\d+) ?(.*)?')
    for line in file_iter:
        # the line ending seems to be inconsistent, so strip all whitespace at the end before doing anything
        line = line.strip()
        if not line:
            continue

        verse_match = VERSE_LINE_REGEX.fullmatch(line)
        if verse_match:
            check_verse_boundary()
            current_verse.append_verse(verse_match[2])
            if is_target():
                current_verse.ref = verse_match[1]

        elif line.startswith(footnote_word):
            # This is a footnote for the target language. Currently only the
            # target language has footnotes on a separate line
            current_verse.append_footnote(line)

        else:
            # Anything else is assumed to be a title line for the next language
            check_verse_boundary()
            current_verse.append_title(line)
    
    # There may be a last unpushed verse at the end
    check_verse_boundary()

    print(f'Retrieved {len(verses)} verses from "{params[PARAM_INPUT_PATH]}"')
    return verses


def export_table_document(verses, params):
    print('Creating Word document with table...')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    # Set orientation to landscape
    section = doc.sections[-1]
    old_width, old_height = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = old_width
    section.page_width = old_height

    # Set the margins
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Add the table
    (col_names, col_widths) = calc_columns(params)
    table = doc.add_table(rows=1, cols=len(col_names), style='Table Grid')
    header_row = table.rows[0]
    for (i, col_name) in enumerate(col_names):
        add_header(col_name, header_row, i)

    for verse in verses:
        add_verse_rows(verse, table, params[PARAM_SPLIT_SENTENCES])

    # Set the column widths. Each cell needs to be set individually
    for idx, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[idx]
            
    return save_document(doc, params[PARAM_OUTPUT_PATH])


def export_simple_document(file_iter, params):
    print('Creating simple Word document...')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    for line in file_iter:
        doc.add_paragraph(text=line)

    return save_document(doc, params[PARAM_OUTPUT_PATH])


def save_document(doc, path:Path):
    try:
        doc.save(str(path))
        return True
    except PermissionError:
        show_error(f'"{path.name}" is currently open. Please close and try again.')
        return False


def calc_columns(params):
    languages = params['languages']

    col_names = ['Verse']
    for lang in languages:
        col_names.append(lang)
    if params[PARAM_NOTES_COLUMN]:
        col_names.append('Notes')

    col_widths = [Cm(3)]
    if len(col_names) == 2:
        col_widths.append(Cm(15))
    elif len(col_names) == 3:
        col_widths.extend([Cm(10), Cm(10)])
    elif len(col_names) == 4 and len(languages) == 3:
        col_widths.extend([Cm(7), Cm(7), Cm(7)])
    elif len(col_names) == 4 and len(languages) == 2:
        col_widths.extend([Cm(8), Cm(8), Cm(5)])
    elif len(col_names) == 5:
        col_widths = [Cm(2.5), Cm(6), Cm(6), Cm(6), Cm(3)]
        
    return (col_names, col_widths)


def add_header(text, row, idx):
    row.cells[idx].text = text
    row.cells[idx].paragraphs[0].runs[0].font.bold = True


SENTENCE_REGEX = re.compile(r'([^.?!]+[.?!]\S*) ?')
def add_verse_rows(verse, table, split_sentences):
    # Always include the whole verse as the first row
    main_row = table.add_row()
    main_row.cells[0].text = verse.ref
    for (i, lang) in enumerate(verse.languages):
        main_row.cells[i+1].text = lang

    if split_sentences:
        lang_sentences = [SENTENCE_REGEX.findall(lang) for lang in verse.languages]

        # Some sentences in either language may be combined so it may
        # not correspond exactly, but an empty line should notify the
        # user of the issue and not mess up the sentence alignment
        num_lines = max([len(s) for s in lang_sentences])
        for sentences in lang_sentences:
            sentences.extend([''] * (num_lines - len(sentences)))

        # Add a row for each sentence
        for i in range(num_lines):
            row = table.add_row()
            row.cells[0].text = f'{verse.ref.strip(".")}:{i+1}'
            for (j, lang) in enumerate(lang_sentences):
                row.cells[j+1].text = lang[i]


def show_error(text):
    print("Error: " + text)
    import ctypes  
    ctypes.windll.user32.MessageBoxW(0, text, "Error Creating Word Document", 0 + 16)


if __name__ == "__main__":
    params = get_params()
    if params:

        def import_and_export(encoding):
            (lines, verses) = import_file(encoding, params)
            if lines:
                if export_simple_document(lines, params):
                    print(f'Successfully exported "{params["output_path"]}"')
            elif verses:
                if export_table_document(verses, params):
                    print(f'Successfully exported "{params["output_path"]}"')

        import_and_export('utf-8-sig')
