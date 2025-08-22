import sys
import re
from pathlib import Path
from difflib import SequenceMatcher

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, RGBColor


# Parameter Name constants
PARAM_INPUT_PATH = 'input_path'
PARAM_OUTPUT_PATH = 'output_path'
PARAM_SPLIT_SENTENCES = 'split_sentences'
PARAM_NOTES_COLUMN = 'add_notes_column'
PARAM_COMPARE = 'compare'
PARAM_TEST = 'test'

# Verse Fields
VERSE_REF = 'ref'
VERSE_TEXT = 'text'


def get_params():
    # usage is: tbta_export_to_table.exe -s -n -c -t "text_file.txt"
    # The text file path is required
    do_split = '-S' in sys.argv or '-s' in sys.argv
    do_notes = '-N' in sys.argv or '-n' in sys.argv
    do_compare = '-C' in sys.argv or '-c' in sys.argv
    is_test = '-T' in sys.argv or '-t' in sys.argv

    non_flag_args = [a for a in sys.argv if not a.startswith('-')]

    if len(non_flag_args) < 2:
        show_error('Please specify a .txt file to import')
        return None

    file_name = non_flag_args[-1]
    file_path = Path(file_name).with_suffix('.txt')
    if not file_path.exists():
        show_error(f'Specified File "{file_name}" does not exist...')
        return None

    return {
        PARAM_INPUT_PATH: file_path,
        PARAM_OUTPUT_PATH: file_path.with_name(f'{file_path.stem}.docx'),
        PARAM_SPLIT_SENTENCES: do_split,
        PARAM_NOTES_COLUMN: do_notes,
        PARAM_COMPARE: do_compare,
        PARAM_TEST: is_test,
    }


def import_text(input_path):
    verses = []
    language_names = []

    VERSE_REF_REGEX = re.compile(r'.+? [\d:]+')
    VERSE_TEXT_REGEX = re.compile(r'(.+?):(.*)')

    print(f'Importing text from "{input_path}"')

    # TODO handle utf-16-le again?
    with input_path.open(encoding='utf-8-sig', newline='\n') as file:
        for line in file:
            # The line ending seems to be inconsistent, so strip all whitespace at the end before doing anything
            line = line.strip()
            if not line:
                # a blank line separates each verse
                continue
            
            ref_match = VERSE_REF_REGEX.fullmatch(line)
            text_match = VERSE_TEXT_REGEX.fullmatch(line)
            if ref_match:
                verses.append({ VERSE_REF: ref_match[0], VERSE_TEXT: {} })

            elif text_match:
                lang_name = text_match[1]
                lang_text = text_match[2]
                verses[-1][VERSE_TEXT][lang_name] = lang_text.strip()

                # set the language names (all verses will have the same languages, in the same order)
                if len(verses) == 1:
                    language_names.append(lang_name)

    print(f'Retrieved {len(verses)} verses')
    return (verses, language_names)


SENTENCE_REGEX = re.compile(r'([^.?!]+[.?!]\S*) ?')
def split_verse_sentences(verses):
    for full_verse in verses:
        yield full_verse
        sentences_by_lang = {lang: SENTENCE_REGEX.findall(text) for lang, text in full_verse[VERSE_TEXT].items()}

        # Some sentences in either language may be combined so it may
        # not correspond exactly, but an empty line should notify the
        # user of the issue and not mess up the sentence alignment
        num_lines = max([len(sentences) for _, sentences in sentences_by_lang.items()])
        for _, sentences in sentences_by_lang.items():
            sentences.extend([''] * (num_lines - len(sentences)))

        # Add a row for each sentence
        for line_num in range(num_lines):
            yield {
                VERSE_REF: f'{full_verse[VERSE_REF]}:{line_num+1}',
                VERSE_TEXT: {lang: sentences[line_num] for lang, sentences in sentences_by_lang.items()},
            }


def export_table(verses, language_names, params):
    print(f'Creating Word document with table rows...')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    # Set orientation to landscape
    section = doc.sections[-1]
    old_width, old_height = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = old_width
    section.page_width = old_height

    # Set the margins
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Add the table
    table = doc.add_table(rows=1, cols=(len(language_names) + 1), style='Table Grid')

    # Add rows for each verse
    for verse in verses:
        verse_row = table.add_row()
        verse_row.cells[0].text = verse[VERSE_REF]

        if params[PARAM_COMPARE]:
            # Compare the last two texts
            *other, old, new = language_names
            for (i, lang_name) in enumerate(other):
                verse_row.cells[i+1].text = verse[VERSE_TEXT][lang_name] or ''

            old_runs, new_runs = compare_text(verse[VERSE_TEXT][old], verse[VERSE_TEXT][new])
            format_text_runs(old_runs, verse_row.cells[-2])
            format_text_runs(new_runs, verse_row.cells[-1])
        else:
            for (i, lang_name) in enumerate(language_names):
                verse_row.cells[i+1].text = verse[VERSE_TEXT][lang_name] or ''

    (col_names, col_widths) = calculate_columns(language_names, params)
    if params[PARAM_NOTES_COLUMN]:
        table.add_column(col_widths[-1])

    # Add the headers
    header_row = table.rows[0]
    for (i, col_name) in enumerate(col_names):
        header_row.cells[i].text = col_name
        header_row.cells[i].paragraphs[0].runs[0].font.bold = True

    # Set the column widths. Each cell needs to be set individually
    for idx, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[idx]

    return save_document(doc, params[PARAM_OUTPUT_PATH])


def compare_text(old, new):
    old_runs, old_i = [], 0
    new_runs, new_i = [], 0

    # Refer to https://docs.python.org/3/library/difflib.html#difflib.SequenceMatcher.get_matching_blocks

    for (a, b, n) in SequenceMatcher(None, old, new).get_matching_blocks():
        if a > old_i:
            old_runs.append({ 'text': old[old_i:a], 'is_diff': True })
        old_runs.append({ 'text': old[a:a+n], 'is_diff': False })
        old_i = a + n

        if b > new_i:
            new_runs.append({ 'text': new[new_i:b], 'is_diff': True })
        new_runs.append({ 'text': new[b:b+n], 'is_diff': False })
        new_i = b + n
    
    return (old_runs, new_runs)


def format_text_runs(runs, table_cell):
    paragraph = table_cell.paragraphs[0]

    for run_data in runs:
        run = paragraph.add_run(text=run_data['text'])
        if run_data['is_diff']:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def calculate_columns(languages, params):
    col_names = ['Verse']
    for lang in languages:
        col_names.append(lang)
    if params[PARAM_NOTES_COLUMN]:
        col_names.append('Notes')

    # Set column widths depending on how many columns there are
    col_widths = [Cm(3)]
    if len(col_names) == 2:
        col_widths.append(Cm(15))
    elif len(col_names) == 3:
        col_widths.extend([Cm(10), Cm(10)])
    elif len(col_names) == 4 and len(languages) == 3:
        col_widths.extend([Cm(7), Cm(7), Cm(7)])
    elif len(col_names) == 4 and len(languages) == 2:
        col_widths.extend([Cm(8.5), Cm(8.5), Cm(4)])
    elif len(col_names) == 5:
        col_widths = [Cm(2.5), Cm(6), Cm(6), Cm(6), Cm(3)]

    return (col_names, col_widths)


def save_document(doc, path:Path):
    try:
        doc.save(str(path))
        print(f'Successfully exported "{path}"')
        return True
    except PermissionError:
        show_error(f'"{path.name}" is currently open. Please close and try again.')
        return False


def show_error(text):
    print("Error: " + text)
    import ctypes  
    ctypes.windll.user32.MessageBoxW(0, text, "Error Creating Word Document", 0 + 16)


if __name__ == "__main__":
    params = get_params()
    if params:
        verses, language_names = import_text(params[PARAM_INPUT_PATH])
        if params[PARAM_SPLIT_SENTENCES]:
            verses = split_verse_sentences(verses)
        if export_table(verses, language_names, params) and not params[PARAM_TEST]:
            print(f'Deleting {params[PARAM_INPUT_PATH]}')
            params[PARAM_INPUT_PATH].unlink()   # delete the original text file
