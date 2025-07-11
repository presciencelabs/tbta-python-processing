import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Parameter Name constants
PARAM_INPUT_PATH = 'input_path'
PARAM_OUTPUT_PATH = 'output_path'
PARAM_TEST = 'test'

def get_params():
    # usage is: tbta_export_to_word.exe (-t) "text_file.txt"
    is_test = '-T' in sys.argv or '-t' in sys.argv
    
    non_flag_args = [a for a in sys.argv if not a.startswith('-')]

    # The text file path is required
    if len(non_flag_args) < 2:
        show_error('Please specify a .txt file to import')
        return None

    file_name = sys.argv[-1]
    file_path = Path(file_name).with_suffix('.txt')
    if not file_path.exists():
        show_error(f'Specified File "{file_name}" does not exist...')
        return None

    return {
        PARAM_INPUT_PATH: file_path,
        PARAM_OUTPUT_PATH: file_path.with_name(f'{file_path.stem}.docx'),
        PARAM_TEST: is_test,
    }


def export_text(params):
    print(f'Creating Word document from "{params[PARAM_INPUT_PATH]}"...')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    # TODO handle utf-16-le again?
    with params[PARAM_INPUT_PATH].open(encoding='utf-8-sig', newline='\n') as file:
        for line in file:
            p = doc.add_paragraph()
            format_text(line, p)

    try:
        doc.save(str(params[PARAM_OUTPUT_PATH]))
        print(f'Successfully exported "{params[PARAM_OUTPUT_PATH]}"')
        return True
    except PermissionError:
        show_error(f'"{params[PARAM_OUTPUT_PATH].name}" is currently open. Please close and try again.')
        return False


def format_text(text, paragraph):
    # Split the text into runs based on asterisks
    runs = [{ 'text': t, 'highlight': i % 2 == 1 } for i, t in enumerate(text.strip().split('*'))]

    # Format the resulting runs
    for run_data in runs:
        run = paragraph.add_run(text=run_data['text'])
        if run_data['highlight']:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def show_error(text):
    print("Error: " + text)
    import ctypes  
    ctypes.windll.user32.MessageBoxW(0, text, "Error Creating Word Document", 0 + 16)


if __name__ == "__main__":
    params = get_params()
    if params:
        if export_text(params) and not params[PARAM_TEST]:
            print(f'Deleting {params[PARAM_INPUT_PATH]}')
            params[PARAM_INPUT_PATH].unlink()   # delete the original text file
